"""Generates the committed .docx fixture corpus for the ooxml/export.js test suite.

Run: pip install -r requirements.txt && python3 fixtures/generate.py

python-docx's high-level API covers paragraphs/headings/lists/tables/runs/comments/images.
Tracked changes, fields, inline content controls, and comment threading/resolution have no
high-level API, so those fixtures drop to python-docx's oxml layer (docx.oxml.OxmlElement +
qn) to build the exact OOXML the exporter needs to see. Comment threading additionally needs
a commentsExtended.xml part that python-docx never writes, so add_comment_threading()
post-processes the saved .docx's zip directly (lxml) to add it, matching what Word itself
would produce (paraId/paraIdParent/done wired through content-types + rels).

Each fixture is documented at its call site below with what it exercises and why. This
script is a generator, not a build step: regenerate and re-commit whenever a fixture's
construction changes.
"""

import base64
import os
import zipfile

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

FIXTURES_DIR = os.path.dirname(os.path.abspath(__file__))

# A minimal 1x1 transparent PNG, used by the images fixture so the corpus doesn't need a
# binary asset checked in separately.
_PNG_1x1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)

_rev_id = [0]


def _next_rev_id():
    _rev_id[0] += 1
    return _rev_id[0]


def add_tracked_insertion(paragraph, text, author="Reviewer A", date="2026-01-01T00:00:00Z"):
    """Wraps a new run in <w:ins>, exercising the exporter's {++...++} path."""
    run = paragraph.add_run(text)
    r_el = run._r
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(_next_rev_id()))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    r_el.addprevious(ins)
    ins.append(r_el)
    return run


def add_tracked_deletion(paragraph, text, author="Reviewer A", date="2026-01-01T00:00:00Z"):
    """Wraps a new run in <w:del> with w:t renamed to w:delText, for {--...--}."""
    run = paragraph.add_run(text)
    r_el = run._r
    for t in r_el.findall(qn("w:t")):
        t.tag = qn("w:delText")
    del_el = OxmlElement("w:del")
    del_el.set(qn("w:id"), str(_next_rev_id()))
    del_el.set(qn("w:author"), author)
    del_el.set(qn("w:date"), date)
    r_el.addprevious(del_el)
    del_el.append(r_el)
    return run


def add_hyperlink(paragraph, text, url):
    """Builds a <w:hyperlink> run pointing at an external relationship (single-run label,
    so the exporter's source map can resolve it -- see export.js's hyperlink handling)."""
    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def add_simple_field(paragraph, instr, cached_text):
    """<w:fldSimple>, for the exporter's locked ⟦field: ...⟧ placeholder."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = cached_text
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def add_list_numbering(paragraph, num_id=1, ilvl=0):
    """Sets direct w:numPr on the paragraph. python-docx's style="List Bullet"/"List
    Number" only sets w:pStyle -- the numbering lives on the STYLE definition, not the
    paragraph -- but the exporter (matching the reference impl) only ever looks at a
    paragraph's own direct numPr, exactly like most real authored Word documents (the
    "Bullets" toolbar button applies numPr directly) do. This numId isn't required to
    resolve to a defined abstractNum for our exporter's purposes."""
    pPr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    pPr.append(num_pr)


def add_inline_content_control(paragraph, text, tag="AutoReviewer.Field"):
    """Inline <w:sdt><w:sdtContent> wrapping a run, for the locked content-control path."""
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    tag_el = OxmlElement("w:tag")
    tag_el.set(qn("w:val"), tag)
    sdt_pr.append(tag_el)
    sdt.append(sdt_pr)
    sdt_content = OxmlElement("w:sdtContent")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    sdt_content.append(r)
    sdt.append(sdt_content)
    paragraph._p.append(sdt)


W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def add_comment_threading(path, thread_spec):
    """Post-processes a saved .docx to add word/commentsExtended.xml (paraId/paraIdParent/
    done), which python-docx's add_comment() never writes. thread_spec maps each comment's
    w:id (as a string, in document order starting at "0") to {"parent": <other w:id>,
    "done": bool}."""
    with zipfile.ZipFile(path) as zin:
        items = [(i.filename, zin.read(i.filename)) for i in zin.infolist() if i.filename != "word/commentsExtended.xml"]
    by_name = dict(items)

    comments_root = etree.fromstring(by_name["word/comments.xml"])
    para_ids = {}
    for comment_el in comments_root.findall(f"{{{W_NS}}}comment"):
        cid = comment_el.get(f"{{{W_NS}}}id")
        p_el = comment_el.find(f"{{{W_NS}}}p")
        pid = format(0x10000000 + int(cid), "08X")
        p_el.set(f"{{{W14_NS}}}paraId", pid)
        para_ids[cid] = pid
        comment_el.set(f"{{{W_NS}}}date", "2026-07-02T00:00:00Z")
    comments_xml = etree.tostring(comments_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    ext_root = etree.Element(f"{{{W15_NS}}}commentsEx", nsmap={"w15": W15_NS})
    for cid, pid in para_ids.items():
        spec = thread_spec.get(cid, {})
        entry = etree.SubElement(ext_root, f"{{{W15_NS}}}commentEx")
        entry.set(f"{{{W15_NS}}}paraId", pid)
        entry.set(f"{{{W15_NS}}}done", "1" if spec.get("done") else "0")
        if "parent" in spec:
            entry.set(f"{{{W15_NS}}}paraIdParent", para_ids[spec["parent"]])
    commentsExtended_xml = etree.tostring(ext_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    rels_root = etree.fromstring(by_name["word/_rels/document.xml.rels"])
    already_related = any(r.get("Target") == "commentsExtended.xml" for r in rels_root)
    if not already_related:
        existing_nums = [int(r.get("Id")[3:]) for r in rels_root if r.get("Id", "").startswith("rId")]
        next_id = f"rId{max(existing_nums) + 1}"
        rel = etree.SubElement(rels_root, f"{{{PKG_REL_NS}}}Relationship")
        rel.set("Id", next_id)
        rel.set("Type", "http://schemas.microsoft.com/office/2011/relationships/commentsExtended")
        rel.set("Target", "commentsExtended.xml")
    rels_xml = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    ct_root = etree.fromstring(by_name["[Content_Types].xml"])
    already_overridden = any(o.get("PartName") == "/word/commentsExtended.xml" for o in ct_root)
    if not already_overridden:
        override = etree.SubElement(ct_root, f"{{{CT_NS}}}Override")
        override.set("PartName", "/word/commentsExtended.xml")
        override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml")
    ct_xml = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    replacements = {
        "word/comments.xml": comments_xml,
        "word/_rels/document.xml.rels": rels_xml,
        "[Content_Types].xml": ct_xml,
    }
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in items:
            zout.writestr(name, replacements.get(name, data))
        zout.writestr("word/commentsExtended.xml", commentsExtended_xml)


def plain_paragraphs():
    """Baseline fixture: a few ordinary paragraphs, no tracked changes, no comments, no
    special formatting. Proves the export pipeline handles the simplest possible document
    and gives a byte-identical-to-reference-impl baseline to diff every other fixture
    against."""
    doc = Document()
    doc.add_paragraph("This is the first paragraph of a plain document.")
    doc.add_paragraph("This is the second paragraph, with nothing tracked.")
    doc.add_paragraph(
        "A third paragraph exists so table/list neighbors in later fixtures "
        "have a plain-paragraph baseline to compare against."
    )
    doc.save(os.path.join(FIXTURES_DIR, "plain-paragraphs.docx"))


def headings_and_lists():
    """Heading 1-3, a bullet list, and a numbered list -- exercises the exporter's
    '#'-prefix and '- '-bullet synthetic-scaffolding paths and multi-level ilvl indent."""
    doc = Document()
    doc.add_heading("Top-Level Heading", level=1)
    doc.add_paragraph("Intro paragraph under the H1.")
    doc.add_heading("A Subsection", level=2)
    doc.add_paragraph("Intro paragraph under the H2.")
    doc.add_heading("A Sub-subsection", level=3)
    for item in ["First bullet", "Second bullet", "Third bullet"]:
        p = doc.add_paragraph(item, style="List Bullet")
        add_list_numbering(p, num_id=1)
    for item in ["First numbered item", "Second numbered item"]:
        p = doc.add_paragraph(item, style="List Number")
        add_list_numbering(p, num_id=2)
    doc.save(os.path.join(FIXTURES_DIR, "headings-and-lists.docx"))


def tables():
    """A 3x3 table with a header row and one short + one longer cell, so cell-level
    source-map blocks and multi-column padding can both be exercised."""
    doc = Document()
    doc.add_paragraph("A table follows.")
    table = doc.add_table(rows=3, cols=3)
    header = table.rows[0].cells
    header[0].text = "Column A"
    header[1].text = "Column B"
    header[2].text = "Column C"
    table.rows[1].cells[0].text = "Row 1, cell A"
    table.rows[1].cells[1].text = "Row 1, cell B is a bit longer than the others"
    table.rows[1].cells[2].text = "1"
    table.rows[2].cells[0].text = "Row 2, cell A"
    table.rows[2].cells[1].text = "Row 2, cell B"
    table.rows[2].cells[2].text = "2"
    doc.add_paragraph("Text after the table.")
    doc.save(os.path.join(FIXTURES_DIR, "tables.docx"))


def hyperlinks_and_images():
    """A single-run hyperlink (exporter maps it as synthetic/opaque, see export.js) and an
    inline image (exporter emits a locked [image: name] placeholder)."""
    doc = Document()
    p1 = doc.add_paragraph("See the ")
    add_hyperlink(p1, "regulatory docket", "https://example.gov/docket/12345")
    p1.add_run(" for background.")

    p2 = doc.add_paragraph()
    run = p2.add_run()
    img_path = os.path.join(FIXTURES_DIR, "_tmp_pixel.png")
    with open(img_path, "wb") as f:
        f.write(_PNG_1x1)
    run.add_picture(img_path, width=100)
    os.remove(img_path)

    doc.add_paragraph("Text after the image.")
    doc.save(os.path.join(FIXTURES_DIR, "hyperlinks-and-images.docx"))


def bold_italic():
    """Bold, italic, and bold+italic runs adjacent to plain runs in the same paragraph, so
    emphasis-marker synthetic-span boundaries (spec §5.2) get exercised on real doc text."""
    doc = Document()
    p = doc.add_paragraph("Plain text, then ")
    p.add_run("bold text").bold = True
    p.add_run(", then ")
    p.add_run("italic text").italic = True
    p.add_run(", then ")
    r = p.add_run("bold italic text")
    r.bold = True
    r.italic = True
    p.add_run(", then plain again.")
    doc.save(os.path.join(FIXTURES_DIR, "bold-italic.docx"))


def tracked_changes():
    """Pre-existing tracked changes: a standalone insertion, a standalone deletion, and an
    adjacent delete+insert (which the exporter collapses into a {~~old~>new~~} substitution)."""
    doc = Document()
    p1 = doc.add_paragraph("The rule shall apply to ")
    add_tracked_insertion(p1, "all covered carriers")
    p1.add_run(".")

    p2 = doc.add_paragraph("This clause is obsolete and ")
    add_tracked_deletion(p2, "should be removed entirely")
    p2.add_run(".")

    p3 = doc.add_paragraph("The deadline is ")
    add_tracked_deletion(p3, "30 days")
    add_tracked_insertion(p3, "60 days")
    p3.add_run(" after publication.")

    doc.save(os.path.join(FIXTURES_DIR, "tracked-changes.docx"))


def comments_threaded():
    """Two anchored comments: one with an unresolved reply (threading via paraIdParent) and
    one marked resolved (done=1). Exercises buildCommentsData's threading/resolution and
    renderThread's reply-arrow rendering."""
    doc = Document()
    p1 = doc.add_paragraph("This sentence has a discussion thread attached to it.")
    doc.add_comment(p1.runs[0], text="Is this the right threshold?", author="Reviewer A")
    p2 = doc.add_paragraph("This sentence has a resolved comment attached to it.")
    doc.add_comment(p2.runs[0], text="Please cite the authority here.", author="Reviewer B")

    path = os.path.join(FIXTURES_DIR, "comments-threaded.docx")
    doc.save(path)
    add_comment_threading(
        path,
        {
            "0": {},
            "1": {"done": True},
        },
    )
    # A reply needs its own <w:comment>, which add_comment() already created for "1" as a
    # sibling top-level comment; re-open, add a genuine reply comment anchored nowhere new
    # (same range as comment 0) so it threads under it via paraIdParent.
    doc2 = Document(path)
    p1_reopened = doc2.paragraphs[0]
    doc2.add_comment(p1_reopened.runs[0], text="Agreed -- flagged for legal review.", author="Reviewer B")
    doc2.save(path)
    add_comment_threading(
        path,
        {
            "0": {},
            "1": {"done": True},
            "2": {"parent": "0"},
        },
    )


def comments_threaded_nested():
    """A deeply threaded discussion thread (up to depth 2) alongside a tracked change in the same paragraph."""
    doc = Document()
    p1 = doc.add_paragraph("This sentence has a discussion thread attached to it.")
    doc.add_comment(p1.runs[0], text="Is this the right threshold?", author="Reviewer A")
    p2 = doc.add_paragraph("This sentence has a resolved comment attached to it.")
    doc.add_comment(p2.runs[0], text="Please cite the authority here.", author="Reviewer B")

    path = os.path.join(FIXTURES_DIR, "comments-threaded-nested.docx")
    doc.save(path)
    add_comment_threading(
        path,
        {
            "0": {},
            "1": {"done": True},
        },
    )
    doc2 = Document(path)
    p1_reopened = doc2.paragraphs[0]
    doc2.add_comment(p1_reopened.runs[0], text="Agreed -- flagged for legal review.", author="Reviewer B")
    doc2.save(path)
    add_comment_threading(
        path,
        {
            "0": {},
            "1": {"done": True},
            "2": {"parent": "0"},
        },
    )
    doc3 = Document(path)
    p1_reopened3 = doc3.paragraphs[0]
    add_tracked_insertion(p1_reopened3, " (with mixed changes)")
    doc3.add_comment(p1_reopened3.runs[0], text="Yes, let's verify.", author="Reviewer C")
    doc3.save(path)
    add_comment_threading(
        path,
        {
            "0": {},
            "1": {"done": True},
            "2": {"parent": "0"},
            "3": {"parent": "2"},
        },
    )


def fields_and_content_controls():
    """A simple field (PAGE, with cached display text) and an inline content control,
    exercising the exporter's locked ⟦field: ...⟧ / content-control placeholders."""
    doc = Document()
    p1 = doc.add_paragraph("Page reference: ")
    add_simple_field(p1, "PAGE", "3")
    p1.add_run(" of the source document.")

    p2 = doc.add_paragraph("Docket number: ")
    add_inline_content_control(p2, "TSA-2026-0042")
    p2.add_run(" (do not edit).")

    doc.save(os.path.join(FIXTURES_DIR, "fields-and-content-controls.docx"))


def stressor():
    """A 300-paragraph stressor (headings + prose + a table every 25 paragraphs) to catch
    perf/determinism regressions on larger documents, per spec §13's "50-page stressor"."""
    doc = Document()
    doc.add_heading("Stress Test Document", level=1)
    for i in range(300):
        if i % 25 == 0:
            doc.add_heading(f"Section {i // 25 + 1}", level=2)
        if i % 50 == 49:
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = f"Metric {i}"
            table.rows[0].cells[1].text = "Value"
            table.rows[1].cells[0].text = "Baseline"
            table.rows[1].cells[1].text = str(i)
            continue
        p = doc.add_paragraph(
            f"Paragraph {i}: this regulatory text discusses cost-benefit considerations, "
            "compliance timelines, and carrier obligations under the proposed rule, "
            "repeated with variation to build up document length for stress testing."
        )
        if i % 10 == 0:
            p.add_run(" ")
            add_tracked_insertion(p, "(flagged for review)")
    doc.save(os.path.join(FIXTURES_DIR, "stressor.docx"))


FIXTURES = [
    plain_paragraphs,
    headings_and_lists,
    tables,
    hyperlinks_and_images,
    bold_italic,
    tracked_changes,
    comments_threaded,
    comments_threaded_nested,
    fields_and_content_controls,
    stressor,
]


def main():
    for fixture in FIXTURES:
        fixture()
        print(f"generated: {fixture.__name__}")


if __name__ == "__main__":
    main()
